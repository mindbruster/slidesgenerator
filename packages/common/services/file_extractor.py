"""
File Extractor Service - extracts text from various file formats
Supports: PDF, DOCX, TXT, MD
"""

import io
from pathlib import Path

from fastapi import HTTPException, UploadFile


# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
ALLOWED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
MIN_TEXT_LENGTH = 50  # Minimum characters for slide generation


class FileExtractorService:
    """Service for extracting text content from uploaded files."""

    @staticmethod
    def validate_file(file: UploadFile) -> str:
        """
        Validate uploaded file.
        Returns the file extension if valid.
        Raises HTTPException if invalid.
        """
        if not file.filename:
            raise HTTPException(status_code=400, detail="Filename is required")

        ext = Path(file.filename).suffix.lower()

        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: {ext}. Allowed: {', '.join(ALLOWED_EXTENSIONS)}",
            )

        return ext

    @staticmethod
    async def read_file_content(file: UploadFile) -> bytes:
        """Read file content with size validation."""
        content = await file.read()

        if len(content) > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=400,
                detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024 * 1024)}MB",
            )

        if len(content) == 0:
            raise HTTPException(status_code=400, detail="File is empty")

        return content

    @staticmethod
    def extract_from_pdf(content: bytes) -> str:
        """Extract text from PDF file."""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="PDF extraction requires PyMuPDF. Install with: pip install pymupdf",
            ) from None

        try:
            doc = fitz.open(stream=content, filetype="pdf")
            text_parts = []

            for page in doc:
                text_parts.append(page.get_text())

            doc.close()
            return "\n".join(text_parts).strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to extract text from PDF: {str(e)}"
            ) from e

    @staticmethod
    def extract_from_docx(content: bytes) -> str:
        """Extract text from DOCX file."""
        try:
            from docx import Document
        except ImportError:
            raise HTTPException(
                status_code=501,
                detail="DOCX extraction requires python-docx. Install with: pip install python-docx",
            ) from None

        try:
            doc = Document(io.BytesIO(content))
            text_parts = []

            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_parts.append(cell.text)

            return "\n".join(text_parts).strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to extract text from DOCX: {str(e)}"
            ) from e

    @staticmethod
    def extract_from_text(content: bytes) -> str:
        """Extract text from TXT/MD file."""
        try:
            # Try UTF-8 first, then fall back to latin-1
            try:
                return content.decode("utf-8").strip()
            except UnicodeDecodeError:
                return content.decode("latin-1").strip()
        except Exception as e:
            raise HTTPException(
                status_code=400, detail=f"Failed to read text file: {str(e)}"
            ) from e

    async def extract(self, file: UploadFile) -> str:
        """
        Extract text from an uploaded file.

        Args:
            file: FastAPI UploadFile object

        Returns:
            Extracted text content

        Raises:
            HTTPException: If file is invalid or extraction fails
        """
        ext = self.validate_file(file)
        content = await self.read_file_content(file)

        if ext == ".pdf":
            text = self.extract_from_pdf(content)
        elif ext == ".docx":
            text = self.extract_from_docx(content)
        elif ext in {".txt", ".md"}:
            text = self.extract_from_text(content)
        else:
            raise HTTPException(status_code=400, detail=f"Unsupported file type: {ext}")

        # Validate extracted text
        if len(text) < MIN_TEXT_LENGTH:
            raise HTTPException(
                status_code=400,
                detail=f"Extracted text is too short ({len(text)} chars). "
                f"Minimum required: {MIN_TEXT_LENGTH} characters.",
            )

        return text


# Singleton instance
file_extractor = FileExtractorService()
